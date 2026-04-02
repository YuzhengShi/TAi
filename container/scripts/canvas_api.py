"""Canvas LMS API helper — reads JSON request from stdin, returns JSON to stdout."""
import json
import re
import sys
import os
import html as html_lib
import io
import subprocess
import tempfile

try:
    import requests
except ImportError:
    print(json.dumps({"error": "requests library not installed"}))
    sys.exit(1)

TOKEN = os.environ.get('CANVAS_API_TOKEN', '')
BASE_URL = os.environ.get('CANVAS_BASE_URL', '').rstrip('/')
COURSE_ID = os.environ.get('CANVAS_COURSE_ID', '')
MAX_PAGES = 5
TIMEOUT = 15

# ── P0: strip raw control chars that break JSON.parse in Node.js ──────────────
_CTRL_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')

def safe_json(text):
    return json.loads(_CTRL_RE.sub('', text))


# ── HTML / text helpers ───────────────────────────────────────────────────────

def strip_html(html):
    text = re.sub(r'<[^>]+>', ' ', html or '')
    text = html_lib.unescape(text)
    return re.sub(r'\s+', ' ', text).strip()

def extract_github_urls(text):
    """Return unique GitHub repo root URLs found in text (github.com and GitHub Enterprise).
    Excludes CDN/asset paths like user-attachments, assets, raw, releases, etc.
    """
    _NON_REPO_OWNERS = {'user-attachments', 'assets', 'raw', 'releases', 'gist', 'orgs', 'users', 'topics'}
    raw = re.findall(r'https?://github[^/\s]*/[\w.\-]+/[\w.\-]+', text or '')
    filtered = [u for u in raw if u.split('/')[-2].lower() not in _NON_REPO_OWNERS]
    return list(dict.fromkeys(filtered))  # dedup, preserve order

def extract_google_docs_urls(text):
    """Extract Google Docs/Sheets/Slides URLs from text (or raw HTML)."""
    raw = re.findall(r'https?://docs\.google\.com/\w+/d/[a-zA-Z0-9_-]+', text or '')
    return list(dict.fromkeys(raw))

MAX_IMAGES_PER_SUBMISSION = 3

def describe_image(image_bytes, filename):
    """Use Bedrock Claude Haiku to describe a screenshot of student work."""
    try:
        import boto3
        import base64

        region = os.environ.get('AWS_REGION', 'us-east-1')
        client = boto3.client('bedrock-runtime', region_name=region)

        ext = filename.rsplit('.', 1)[-1].lower()
        media_type = {
            'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
            'gif': 'image/gif', 'webp': 'image/webp',
        }.get(ext, 'image/png')

        b64 = base64.b64encode(image_bytes).decode()

        resp = client.invoke_model(
            modelId='us.anthropic.claude-haiku-4-5-20251001',
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 400,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                        {"type": "text", "text": "Describe what this screenshot shows about a CS student's work. Focus on architecture, code, test results, deployment output, or system design. 2-3 sentences."}
                    ]
                }]
            })
        )
        body = json.loads(resp['body'].read())
        return body['content'][0]['text']
    except Exception as e:
        return f"[Image analysis unavailable: {e}]"


# ── Attachment download + text extraction ─────────────────────────────────────

_TEXT_EXTS = {
    'md', 'tf', 'tfvars', 'go', 'py', 'java', 'js', 'ts', 'jsx', 'tsx',
    'yaml', 'yml', 'toml', 'sh', 'bash', 'txt', 'csv', 'html', 'htm',
    'json', 'xml', 'dockerfile', 'makefile', 'rb', 'rs', 'cpp', 'c', 'h',
}
_SKIP_EXTS = {'mov', 'mp4', 'avi', 'mkv', 'mp3', 'wav'}
_ARCHIVE_EXTS = {'zip', 'tar', 'gz', 'tgz'}
_IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff'}
MAX_ATTACHMENT_SIZE = 50 * 1024 * 1024  # 50 MB hard cap


def extract_archive(raw, filename):
    """Extract text files from a zip/tar/gz archive. Returns (text, error_note)."""
    import zipfile
    import tarfile

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    parts = []
    skipped = []
    total_text = 0
    MAX_EXTRACTED_TEXT = 500_000  # 500 KB total text budget

    def _should_read(member_name):
        """Check if an archive member is a text file we should read."""
        m_ext = member_name.rsplit('.', 1)[-1].lower() if '.' in member_name else ''
        m_lower = member_name.rsplit('/', 1)[-1].lower() if '/' in member_name else member_name.lower()
        if m_ext in _SKIP_EXTS or m_ext in _IMAGE_EXTS or m_ext in _ARCHIVE_EXTS:
            return False
        if m_ext in _TEXT_EXTS or m_lower in ('dockerfile', 'makefile', 'procfile'):
            return True
        if m_ext == 'pdf':
            return True
        return False

    try:
        if ext == 'zip':
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                for info in sorted(zf.infolist(), key=lambda i: i.filename):
                    if info.is_dir() or info.file_size == 0:
                        continue
                    if info.file_size > 5 * 1024 * 1024:
                        skipped.append(f"{info.filename} (>5MB)")
                        continue
                    if not _should_read(info.filename):
                        skipped.append(info.filename)
                        continue
                    if total_text >= MAX_EXTRACTED_TEXT:
                        skipped.append(f"{info.filename} (text budget exceeded)")
                        continue
                    try:
                        content = zf.read(info.filename)
                        text = content.decode('utf-8', errors='replace')
                        chunk = text[:4000]
                        parts.append(f"=== {info.filename} ===\n{chunk}")
                        total_text += len(chunk)
                    except Exception:
                        skipped.append(f"{info.filename} (read error)")

        elif ext in ('tar', 'gz', 'tgz'):
            mode = 'r:gz' if ext in ('gz', 'tgz') else 'r:'
            with tarfile.open(fileobj=io.BytesIO(raw), mode=mode) as tf:
                for member in sorted(tf.getmembers(), key=lambda m: m.name):
                    if not member.isfile() or member.size == 0:
                        continue
                    if member.size > 5 * 1024 * 1024:
                        skipped.append(f"{member.name} (>5MB)")
                        continue
                    if not _should_read(member.name):
                        skipped.append(member.name)
                        continue
                    if total_text >= MAX_EXTRACTED_TEXT:
                        skipped.append(f"{member.name} (text budget exceeded)")
                        continue
                    try:
                        f = tf.extractfile(member)
                        if f is None:
                            continue
                        content = f.read()
                        text = content.decode('utf-8', errors='replace')
                        chunk = text[:4000]
                        parts.append(f"=== {member.name} ===\n{chunk}")
                        total_text += len(chunk)
                    except Exception:
                        skipped.append(f"{member.name} (read error)")
        else:
            return None, f"[{filename}: unsupported archive format]"

    except (zipfile.BadZipFile, tarfile.TarError) as e:
        return None, f"[{filename}: corrupt archive — {e}]"
    except Exception as e:
        return None, f"[{filename}: archive extraction failed — {e}]"

    if not parts:
        note = f"[{filename}: no readable text files found]"
        if skipped:
            note += f" Skipped: {', '.join(skipped[:10])}"
        return None, note

    text = "\n\n".join(parts)
    note = None
    if skipped:
        note = f"[Skipped {len(skipped)} non-text files in {filename}]"
    return text, note


def download_file(url):
    """Download from Canvas URL. Try with auth first (for Canvas file URLs),
    then without (for pre-signed S3 URLs)."""
    try:
        resp = requests.get(url, headers={"Authorization": f"Bearer {TOKEN}"}, timeout=30)
        if resp.status_code == 403:
            # Pre-signed URL — no auth header
            resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        return resp.content
    except Exception as e:
        raise RuntimeError(f"Download failed: {e}")


def read_attachment(url, filename, content_type, size=0):
    """Download and extract text from a Canvas attachment.
    Returns (text_or_None, error_note_or_None)."""
    if not url:
        return None, f"No URL for {filename}"

    ct = (content_type or '').lower()
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    fname_lower = filename.lower()

    # Skip large / media files before downloading
    if size > MAX_ATTACHMENT_SIZE:
        return None, f"[SKIPPED: {filename} too large ({size // 1024 // 1024} MB)]"
    if ext in _SKIP_EXTS:
        return None, f"[SKIPPED: {filename} (binary/media)]"
    if ext in _ARCHIVE_EXTS:
        try:
            raw = download_file(url)
        except RuntimeError as e:
            return None, str(e)
        return extract_archive(raw, filename)
    if ext in _IMAGE_EXTS:
        if not os.environ.get('AWS_ACCESS_KEY_ID'):
            return None, f"[IMAGE: {filename} — no AWS credentials for vision]"
        if size > 5 * 1024 * 1024:
            return None, f"[IMAGE: {filename} — too large for vision]"
        try:
            raw = download_file(url)
            desc = describe_image(raw, filename)
            return f"[Screenshot: {filename}]\n{desc}", None
        except Exception as e:
            return None, f"[IMAGE: {filename} — {e}]"

    try:
        raw = download_file(url)
    except RuntimeError as e:
        return None, str(e)

    # ── PDF ──────────────────────────────────────────────────────────────────
    if 'pdf' in ct or ext == 'pdf':
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(raw)
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ['pdftotext', '-layout', tmp_path, '-'],
                capture_output=True, text=True, timeout=30
            )
            text = result.stdout.strip()
            return (text if text else None,
                    None if text else f"[PDF {filename}: image-based, no extractable text]")
        except Exception as e:
            return None, f"[PDF {filename}: pdftotext failed — {e}]"
        finally:
            try: os.unlink(tmp_path)
            except: pass

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if ext == 'docx' or 'wordprocessingml' in ct:
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)
            # Also grab table cells
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        parts.append(' | '.join(row_texts))
            text = '\n'.join(parts)
            return (text if text else None,
                    None if text else f"[DOCX {filename}: no text content]")
        except ImportError:
            return None, f"[DOCX {filename}: python-docx not installed]"
        except Exception as e:
            return None, f"[DOCX {filename}: parse error — {e}]"

    # ── Plain text variants ───────────────────────────────────────────────────
    is_text_ct = any(ct.startswith(p) for p in ('text/', 'application/json', 'application/xml'))
    is_text_ext = ext in _TEXT_EXTS or fname_lower in ('dockerfile', 'makefile', 'procfile')
    if is_text_ct or is_text_ext or 'octet-stream' in ct:
        # octet-stream is used for .tf / .tfvars — try to decode as text
        try:
            return raw.decode('utf-8', errors='replace'), None
        except Exception as e:
            return None, f"[{filename}: decode error — {e}]"

    return None, f"[{filename} ({ct}): unsupported format]"


# ── Core API helpers ──────────────────────────────────────────────────────────

def api_get(endpoint, params=None):
    """GET with pagination, returns combined list or single object."""
    url = f"{BASE_URL}{endpoint}"
    headers = {"Authorization": f"Bearer {TOKEN}"}
    all_results = []
    pages = 0

    while url and pages < MAX_PAGES:
        resp = requests.get(url, headers=headers, params=params, timeout=TIMEOUT)
        resp.raise_for_status()
        data = safe_json(resp.text)          # P0: use safe_json
        if isinstance(data, list):
            all_results.extend(data)
        else:
            return data
        url = None
        params = None
        link = resp.headers.get('Link', '')
        for part in link.split(','):
            if 'rel="next"' in part:
                url = part.split('<')[1].split('>')[0]
        pages += 1

    return all_results


def fmt_assignment(a):
    return {
        "id": a.get("id"),
        "name": a.get("name"),
        "due_at": a.get("due_at"),
        "points_possible": a.get("points_possible"),
        "submission_types": a.get("submission_types"),
        "html_url": a.get("html_url"),
        "published": a.get("published"),
    }


def fmt_submission(s):
    user = s.get("user", {}) or {}
    return {
        "user_id": s.get("user_id"),
        "user_name": user.get("name") or user.get("sortable_name"),
        "submitted_at": s.get("submitted_at"),
        "late": s.get("late"),
        "missing": s.get("missing"),
        "score": s.get("score"),
        "grade": s.get("grade"),
        "workflow_state": s.get("workflow_state"),
    }


# ── P1: rich student submission content extraction ────────────────────────────

def extract_submission_content(sub):
    """Extract readable text from a Canvas submission object.

    Returns a dict with:
      submission_type, full_text, github_urls, attachment_notes
    """
    submission_type = sub.get("submission_type", "") or ""
    full_text = ""
    github_urls = []
    google_docs_urls = []
    attachment_notes = []

    if submission_type == "online_text_entry":
        body_html = sub.get("body", "") or ""
        github_urls = extract_github_urls(body_html)        # raw HTML catches href URLs
        google_docs_urls = extract_google_docs_urls(body_html)
        full_text = strip_html(body_html)

    elif submission_type == "online_upload":
        attachments = sub.get("attachments", []) or []
        parts = []
        images_analyzed = 0
        for att in attachments:
            filename = att.get("filename") or att.get("display_name") or "unknown"
            content_type = att.get("content-type") or att.get("mime_class") or ""
            size = att.get("size") or 0
            url = att.get("url") or att.get("download_url") or ""
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

            # Cap image vision calls per submission
            if ext in _IMAGE_EXTS and images_analyzed >= MAX_IMAGES_PER_SUBMISSION:
                attachment_notes.append(f"[SKIPPED: {filename} — image limit ({MAX_IMAGES_PER_SUBMISSION})]")
                continue

            text, note = read_attachment(url, filename, content_type, size)
            if text:
                if ext in _IMAGE_EXTS:
                    images_analyzed += 1
                github_urls.extend(extract_github_urls(text))
                parts.append(f"=== {filename} ===\n{text[:4000]}")
            if note:
                attachment_notes.append(note)

        full_text = "\n\n".join(parts)
        github_urls = list(dict.fromkeys(github_urls))  # dedup
        google_docs_urls = extract_google_docs_urls(full_text)

    elif submission_type in ("online_url", "basic_lti_launch"):
        url = sub.get("url") or ""
        full_text = f"Submission URL: {url}"
        github_urls = extract_github_urls(url)

    return {
        "submission_type": submission_type,
        "submitted_at": sub.get("submitted_at"),
        "score": sub.get("score"),
        "grade": sub.get("grade"),
        "late": sub.get("late"),
        "full_text": full_text,
        "github_urls": github_urls,
        "google_docs_urls": google_docs_urls,
        "attachment_notes": attachment_notes,
    }


# ── Action dispatcher ─────────────────────────────────────────────────────────

def handle(req):
    action = req.get("action", "")
    params = req.get("params", {})
    c = COURSE_ID

    if action == "assignments":
        data = api_get(f"/courses/{c}/assignments", {"per_page": 50, "order_by": "due_at"})
        return [fmt_assignment(a) for a in data]

    elif action == "assignment_detail":
        aid = params.get("assignment_id", "")
        if not aid:
            return {"error": "assignment_id required"}
        data = api_get(f"/courses/{c}/assignments/{aid}", {"include[]": "rubric"})
        # Extract rubric if present, attach as top-level fields
        if data.get("rubric"):
            data["rubric_criteria"] = [
                {
                    "id": cr.get("id"),
                    "description": cr.get("description"),
                    "long_description": cr.get("long_description"),
                    "points": cr.get("points"),
                    "ratings": [{"description": r.get("description"), "points": r.get("points")} for r in (cr.get("ratings") or [])],
                }
                for cr in data["rubric"]
            ]
        return data

    elif action == "submissions":
        aid = params.get("assignment_id", "")
        if not aid:
            return {"error": "assignment_id required"}
        data = api_get(f"/courses/{c}/assignments/{aid}/submissions", {"per_page": 50, "include[]": "user"})
        return [fmt_submission(s) for s in data]

    elif action == "student_submission":
        aid = params.get("assignment_id", "")
        uid = params.get("user_id", "")
        if not aid or not uid:
            return {"error": "assignment_id and user_id required"}
        sub = api_get(f"/courses/{c}/assignments/{aid}/submissions/{uid}",
                      {"include[]": ["submission_history", "submission_comments", "rubric_assessment"]})
        result = extract_submission_content(sub)
        # Also scan submission comments for GitHub URLs (students often paste repo link in comments)
        for comment in (sub.get("submission_comments") or []):
            comment_text = comment.get("comment") or ""
            for url in extract_github_urls(comment_text):
                if url not in result["github_urls"]:
                    result["github_urls"].append(url)
        result["rubric_assessment"] = sub.get("rubric_assessment") or {}
        return result

    elif action == "grades":
        data = api_get(f"/courses/{c}/students/submissions", {"per_page": 50, "student_ids[]": "all"})
        return data

    elif action == "my_grades":
        uid = params.get("user_id", "")
        if not uid:
            return {"error": "user_id required (bound automatically for students)"}
        # Fetch all submissions for this single student
        subs = api_get(f"/courses/{c}/assignments", {"per_page": 50, "order_by": "due_at"})
        results = []
        for a in subs:
            aid = a.get("id")
            try:
                sub = api_get(f"/courses/{c}/assignments/{aid}/submissions/{uid}")
                results.append({
                    "assignment": a.get("name"),
                    "due_at": a.get("due_at"),
                    "points_possible": a.get("points_possible"),
                    "score": sub.get("score"),
                    "grade": sub.get("grade"),
                    "late": sub.get("late"),
                    "missing": sub.get("missing"),
                    "submitted_at": sub.get("submitted_at"),
                })
            except Exception:
                results.append({
                    "assignment": a.get("name"),
                    "due_at": a.get("due_at"),
                    "points_possible": a.get("points_possible"),
                    "score": None,
                    "grade": None,
                    "error": "Could not fetch submission",
                })
        return results

    elif action == "announcements":
        data = api_get(f"/courses/{c}/discussion_topics", {"only_announcements": "true", "per_page": 20})
        return [{"id": a["id"], "title": a.get("title"), "message": strip_html(a.get("message") or "")[:500], "posted_at": a.get("posted_at")} for a in data]

    elif action == "discussions":
        data = api_get(f"/courses/{c}/discussion_topics", {"per_page": 20})
        return [{"id": d["id"], "title": d.get("title"), "message": strip_html(d.get("message") or "")[:300], "posted_at": d.get("posted_at")} for d in data]

    elif action == "discussion_detail":
        tid = params.get("topic_id", "")
        if not tid:
            return {"error": "topic_id required"}
        return api_get(f"/courses/{c}/discussion_topics/{tid}")

    elif action == "upcoming":
        data = api_get(f"/courses/{c}/assignments", {"bucket": "upcoming", "per_page": 20, "order_by": "due_at"})
        return [fmt_assignment(a) for a in data]

    elif action == "modules":
        data = api_get(f"/courses/{c}/modules", {"include[]": "items", "per_page": 50})
        return [{"id": m["id"], "name": m.get("name"), "position": m.get("position"), "items_count": m.get("items_count"), "items": [{"title": i.get("title"), "type": i.get("type"), "content_id": i.get("content_id"), "external_url": i.get("external_url"), "html_url": i.get("html_url"), "page_url": i.get("page_url")} for i in (m.get("items") or [])[:20]]} for m in data]

    elif action == "module_items":
        mid = params.get("module_id", "")
        if not mid:
            return {"error": "module_id required"}
        data = api_get(f"/courses/{c}/modules/{mid}/items", {"per_page": 50})
        return [{"id": i.get("id"), "title": i.get("title"), "type": i.get("type"), "content_id": i.get("content_id"), "external_url": i.get("external_url"), "html_url": i.get("html_url"), "page_url": i.get("page_url")} for i in data]

    elif action == "users":
        data = api_get(f"/courses/{c}/users", {"enrollment_type": "student", "per_page": 50})
        return [{"id": u["id"], "name": u.get("name"), "sortable_name": u.get("sortable_name"), "email": u.get("email")} for u in data]

    elif action == "pages":
        data = api_get(f"/courses/{c}/pages", {"per_page": 50, "sort": "title"})
        return [{"url": p.get("url"), "title": p.get("title"), "updated_at": p.get("updated_at"), "published": p.get("published")} for p in data]

    elif action == "page_detail":
        page_url = params.get("page_url", "")
        if not page_url:
            return {"error": "page_url required (the url slug, not full URL)"}
        data = api_get(f"/courses/{c}/pages/{page_url}")
        body_text = strip_html(data.get("body") or "")
        return {"url": data.get("url"), "title": data.get("title"), "body": body_text[:4000], "updated_at": data.get("updated_at")}

    elif action == "syllabus":
        data = api_get(f"/courses/{c}", {"include[]": "syllabus_body"})
        body_text = strip_html(data.get("syllabus_body") or "")
        return {"course_name": data.get("name"), "syllabus": body_text[:6000]}

    elif action == "files":
        data = api_get(f"/courses/{c}/files", {"per_page": 50, "sort": "updated_at", "order": "desc"})
        return [{"id": f.get("id"), "display_name": f.get("display_name"), "filename": f.get("filename"), "size": f.get("size"), "content_type": f.get("content-type"), "updated_at": f.get("updated_at"), "url": f.get("url")} for f in data]

    elif action == "discussion_entries":
        tid = params.get("topic_id", "")
        if not tid:
            return {"error": "topic_id required"}
        data = api_get(f"/courses/{c}/discussion_topics/{tid}/entries", {"per_page": 50})
        return [{"id": e.get("id"), "user_id": e.get("user_id"), "user_name": (e.get("user") or {}).get("display_name"), "message": strip_html(e.get("message") or "")[:500], "created_at": e.get("created_at")} for e in data]

    elif action == "file_content":
        file_id = params.get("file_id", "")
        if not file_id:
            return {"error": "file_id required (get from 'files' action)"}
        file_meta = api_get(f"/files/{file_id}")
        display_name = file_meta.get("display_name", "")
        content_type = file_meta.get("content-type", "")
        size = file_meta.get("size", 0)
        download_url = file_meta.get("url")
        if not download_url:
            return {"error": "No download URL available for this file"}
        text, note = read_attachment(download_url, display_name, content_type, size)
        if text:
            return {"display_name": display_name, "content_type": content_type, "content": text[:6000], "truncated": len(text) > 6000}
        return {"display_name": display_name, "content_type": content_type, "error": note or "Could not extract text"}

    else:
        return {"error": f"Unknown action: {action}", "available": ["assignments", "assignment_detail", "submissions", "student_submission", "grades", "my_grades", "announcements", "discussions", "discussion_detail", "discussion_entries", "upcoming", "modules", "module_items", "pages", "page_detail", "syllabus", "files", "file_content", "users"]}


try:
    raw = sys.stdin.read().strip()
    if not raw:
        print(json.dumps({"error": "No input provided"}))
        sys.exit(1)
    if not TOKEN:
        print(json.dumps({"error": "CANVAS_API_TOKEN not set. Add it to .env"}))
        sys.exit(1)
    if not BASE_URL:
        print(json.dumps({"error": "CANVAS_BASE_URL not set. Add it to .env"}))
        sys.exit(1)
    if not COURSE_ID:
        print(json.dumps({"error": "CANVAS_COURSE_ID not set. Add it to .env"}))
        sys.exit(1)

    req = json.loads(raw)
    result = handle(req)
    print(json.dumps(result, indent=2, default=str))
except requests.exceptions.HTTPError as e:
    print(json.dumps({"error": f"Canvas API error: {e.response.status_code} {e.response.text[:500]}"}))
    sys.exit(1)
except Exception as e:
    print(json.dumps({"error": f"{type(e).__name__}: {e}"}))
    sys.exit(1)
